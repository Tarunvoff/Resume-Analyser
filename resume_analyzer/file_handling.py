import os
import zipfile
import xml.etree.ElementTree as ET
import pdfplumber
import docx

class FileHandler:
    @staticmethod
    def extract_text(file_path):
        import os

        file_extension = os.path.splitext(file_path)[1]
        """Extract text from a DOCX or PDF file."""
        print(f"Processing file: {file_path}")  # Debugging line
        file_extension = os.path.splitext(file_path)[1].lower()
        print(f"Detected extension: {file_extension}")  # Debugging line

        if file_extension == ".docx":   
            return FileHandler.extract_text_from_docx(file_path)
        elif file_extension == ".pdf":
            return FileHandler.extract_text_from_pdf(file_path)
        if not file_extension:
            file_extension = ".pdf"  # or ".docx" based on expected input
            return FileHandler.extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}. Please provide a '.docx' or '.pdf' file.")

    @staticmethod
    def extract_text_from_docx(docx_path):
        """Extract text from a DOCX file, including tables."""
        try:
            doc = docx.Document(docx_path)  # Load DOCX file
            text_list = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_list.append(para.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_list.append(" | ".join(row_text))  # Format table data

            return "\n".join(text_list) if text_list else "No readable text found in the document."
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"

    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """Extract text from a PDF file using pdfplumber, handling missing text."""
        try:
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n"
            return text.strip() if text else "No readable text found in the PDF."
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"

    @staticmethod
    def get_file_page_count(file_path):
        """Get the number of pages in a DOCX or PDF file."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == ".pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    return len(pdf.pages)
            except Exception:
                return "Page count not available (PDF may be corrupted)."
        
        elif file_extension == ".docx":
            try:
                with zipfile.ZipFile(file_path) as docx_zip:
                    xml_content = docx_zip.read("docProps/app.xml")
                    root = ET.fromstring(xml_content)
                    pages = root.find("{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages")
                    return int(pages.text) if pages is not None else "Page count not available"
            except (KeyError, zipfile.BadZipFile):
                return "Page count not available (DOCX may be corrupted)."
        
        else:
            return "Unsupported file type. Please provide a '.docx' or '.pdf' file."
